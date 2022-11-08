# coding: utf-8

import sys
import datetime

from docx import Document

from gycommon import commoninfo

DETAIL_SELLER_ID_INDEX = 0
DETAIL_SELLER_INDEX = 1
DETAIL_BUYER_INDEX = 2
DETAIL_SALE_TIME_INDEX = 3
DETAIL_PRICE_INDEX = 4
DETAIL_TOKEN_ID_INDEX = 5
DETAIL_PROD_ID_INDEX = 6
DETAIL_DETAIL_ID_INDEX = 7
DETAIL_ITEM_COUNT = 8

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("{} <casting_id> <tag> <bucket_size>".format(sys.argv[0]))
        sys.exit(1)
    casting_id = int(sys.argv[1])
    if casting_id not in commoninfo.CastingId2MetaInfo:
        print("CastingId:{} not exist!")
        sys.exit(1)
    casting_name = commoninfo.CastingId2MetaInfo[casting_id][0]
    casting_ch_name = commoninfo.CastingId2MetaInfo[casting_id][1]
    tag = sys.argv[2]
    bucket_size = int(sys.argv[3])

    token2price = {}
    trans_logs_file_name = "data/" + commoninfo.Transaction_Logs_File_Name.format(casting_name)
    with open(trans_logs_file_name, encoding="utf-8-sig") as trans_logs_file:
        # 加载已有的交易信息
        for line in trans_logs_file:
            items = line.strip().split(",")
            #seller_id = int(items[DETAIL_SELLER_ID_INDEX])
            #seller_name = items[DETAIL_SELLER_INDEX]
            sale_time = datetime.datetime.strptime(
                items[DETAIL_SALE_TIME_INDEX], "%Y/%m/%d %H:%M:%S")
            price = float(items[DETAIL_PRICE_INDEX])
            token_id = int(items[DETAIL_TOKEN_ID_INDEX])
            if token_id not in token2price:
                token2price[token_id] = [price, sale_time]
            else:
                if sale_time > token2price[token_id][1]:
                    token2price[token_id] = [price, sale_time]
   
    
     
    # 生成docx文件
    timestamp = datetime.datetime.now().strftime("%Y/%m/%d %H:%M:%S") 
    docx_file_name = "data/{}_{}.docx".format(
        casting_name, tag
    )
    doc = Document(docx_file_name) # 追加到已有文件

    # 按成交价格倒序排列
    token_price_info = []
    for (token_id, (price, _)) in token2price.items():
        token_price_info.append([token_id, price])
    token_price_info.sort(key=lambda i: i[1], reverse=True)
    token_cnt = len(token_price_info)
    if token_cnt == 0:
        sys.exit(0)
    highest_price = token_price_info[0][1]
    bucket_cnt = int(highest_price) // bucket_size + 1
    buckets = [0] * bucket_cnt
    for (_, price) in token_price_info:
        bucket_index = int(price) // bucket_size
        #for i in range(bucket_index+1):
        #    buckets[i] += 1
        buckets[bucket_index] += 1
    result_file_name = "data/_analyze_trans_price_result_{}_{}.csv".format(casting_name, tag)
    with open(result_file_name, "w", encoding="utf-8-sig") as result_file:
        doc.add_heading("{}筹码分布数据".format(casting_ch_name), level=3)
        doc.add_paragraph("截至{}, 总数量:{}".format(timestamp, token_cnt))
        table = doc.add_table(1, 4)
        table.style = "TableGrid"
        heading_cells = table.rows[0].cells
        heading_cells[0].text = "价格区间"
        heading_cells[1].text = "数量"
        heading_cells[2].text = "占比"
        heading_cells[2].text = "累积"
        
        result_file.write("总数量:{}\n".format(token_cnt))
        result_file.write("价格区间,数量,占比,累积\n")
        for i in range(len(buckets)):
            result_file.write("{},{},{:.2%},{:.2%}\n".format(
                "[{}-{})".format(i * bucket_size, (i + 1) * bucket_size),
                buckets[i], buckets[i] / token_cnt,
                sum(buckets[:i+1]) / token_cnt
            ))
            cells = table.add_row().cells
            cells[0].text = "[{}-{})".format(i * bucket_size, (i + 1) * bucket_size)
            cells[1].text = str(buckets[i])
            cells[2].text = "{:.2%}".format(buckets[i] / token_cnt)
            cells[3].text = "{:.2%}".format(sum(buckets[:i+1]) / token_cnt)
   
    doc.save(docx_file_name)
