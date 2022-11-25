# coding: utf-8

import sys
import datetime
import requests
import time
import os

from docx import Document

from gycommon import commoninfo
from gycommon import utils

GET_ON_SALE_LIST_URL = "https://api.gandart.com/market/api/v2/resaleManage/resale/onSale"
PAGE_SIZE = 500
TIME_OUT = 3
GET_PRODUCT_DETAIL_URL = "https://api.gandart.com/market/api/v2/resaleManage/resale/collectionDetails"
GET_TRANS_INFO_URL = "https://api.gandart.com/market/api/v2/resaleManage/resale/transactionInfo"

TRAN_STATUS_SALED = 1

PROD_ID_INDEX = 0
TOKEN_ID_INDEX = 1
PRICE_INDEX = 2
CREATE_TINE_INDEX = 3

DETAIL_SELLER_ID_INDEX = 0
DETAIL_SELLER_INDEX = 1
DETAIL_BUYER_INDEX = 2
DETAIL_SALE_TIME_INDEX = 3
DETAIL_PRICE_INDEX = 4
DETAIL_TOKEN_ID_INDEX = 5
DETAIL_PROD_ID_INDEX = 6
DETAIL_DETAIL_ID_INDEX = 7
DETAIL_ITEM_COUNT = 8

def get_saled_products(casting_id):
    saled_prods = []
    data = {
        "castingId": casting_id,
        "page":1,
        "pageSize":PAGE_SIZE,
        "sort":2,
        "transactionStatus": TRAN_STATUS_SALED,
    }
    data = utils.decorate_api_data(data)
    res = utils.post_requests_json(GET_ON_SALE_LIST_URL, data=data, timeout=TIME_OUT)
    if not res:
        return (1, None)
    if res["code"] != 0:
        return (res["code"], None)
    else:
        for pinfo in res["obj"]["list"]:
            saled_prods.append(
                [pinfo["id"], pinfo["viewSort"], float(pinfo["resalePrice"])])
        total = res["obj"]["total"]
        for pindex in range(2, total // PAGE_SIZE + 2):
            print("{}/{}".format(pindex, total // PAGE_SIZE + 1))
            data = {
                "castingId": casting_id,
                "page":pindex,
                "pageSize":PAGE_SIZE,
                "sort":2,
                "transactionStatus": TRAN_STATUS_SALED,
            }
            res = utils.post_requests_json(GET_ON_SALE_LIST_URL, data=data, timeout=3)
            if res["code"] != 0:
                return (res["code"], None)
            else:
                for pinfo in res["obj"]["list"]:
                    saled_prods.append(
                        [pinfo["id"], pinfo["viewSort"], float(pinfo["resalePrice"])])
            # 防止被封禁
            time.sleep(1) 
    return (0, saled_prods)

def clean_name(name):
    return name.replace(",", "")

def get_product_detail(prod_id):
    detail_info = [None] * DETAIL_ITEM_COUNT

    # Get detail data
    data = {
        "transactionRecordId": prod_id,
    }
    data = utils.decorate_api_data(data)
    detail_id = None
    user_id = None
    created_time = None
    holder_name = None
    while True:
        try: 
            res = requests.post(GET_PRODUCT_DETAIL_URL, data=data, timeout=TIME_OUT).json()
            if res["code"] != 0:
                return None
            else:
                detail_id = res["obj"]["detailId"]
                user_id = res["obj"]["userId"]
                created_time = datetime.datetime.strptime(res["obj"]["created"], "%Y-%m-%d %H:%M:%S")
                holder_name = clean_name(res["obj"]["holderName"])
                #print(detail_id, user_id, created_time)
                break
        except Exception as e:
            time.sleep(3)
            print(e)
    if not detail_id:
        return None
    
    # Get all recent transaction data
    data = {
        "detailId": detail_id,
        "page": 1,
        "pageSize": PAGE_SIZE,
    }
    data = utils.decorate_api_data(data)
    while True:
        try: 
            res = requests.post(GET_TRANS_INFO_URL, data=data, timeout=TIME_OUT).json()
            if res["code"] != 0:
                return None
            else:
                trans = res["obj"]["history"]["list"]
                if len(trans) < 2:
                    print("detailId:{} 未获取到足够的交易信息!".format(detail_id))
                    return None
                # 找到相应的挂单记录，与detail接口的时间一致
                target_sell_index = None
                for i in range(len(trans)):
                    tinfo = trans[i]
                    if created_time == datetime.datetime.strptime(tinfo["created"], "%Y-%m-%d %H:%M:%S") and \
                        tinfo["sellPrice"]:
                        target_sell_index = i
                        break
                if target_sell_index == None:
                    print("productId:{}, detailId:{} 未发现匹配的交易记录!".format(prod_id, detail_id))
                    return None
                
                sell_price = trans[target_sell_index]["sellPrice"]
                target_buy_index = target_sell_index - 1
                
                # 如果买入和挂单时间一样，有时顺序会被打乱
                if target_buy_index < 0 or (not trans[target_buy_index]["buyPrice"]):
                    target_buy_index = target_sell_index + 1
                if target_buy_index >= len(trans): # ProdId:119695, DetailId:97124
                    return None

                buy_price = trans[target_buy_index]["buyPrice"]
                # 有时价格会有微小差别（比如差一分)
                if abs(float(sell_price) - float(buy_price)) > 0.02:
                    print("detailId:{} 买入{}和卖出{}价格不匹配".format(detail_id, buy_price, sell_price))
                    return None
                buyer_name = clean_name(trans[target_buy_index]["nickName"])
                #seller_name = trans[target_sell_index]["nickName"]
                
                detail_info[DETAIL_SELLER_ID_INDEX] = user_id
                detail_info[DETAIL_BUYER_INDEX] = buyer_name
                detail_info[DETAIL_SELLER_INDEX] = holder_name
                detail_info[DETAIL_SALE_TIME_INDEX] = datetime.datetime.strptime(trans[target_buy_index]["created"], "%Y-%m-%d %H:%M:%S")
                detail_info[DETAIL_PRICE_INDEX] = float(buy_price)
                detail_info[DETAIL_PROD_ID_INDEX] = prod_id
                detail_info[DETAIL_DETAIL_ID_INDEX] = detail_id
                break
        except Exception as e:
            time.sleep(2)
            print(e)
    
    return detail_info

if __name__ == "__main__":
    if len(sys.argv) < 4:
        print("{} <casting_id> <date>(YYYYmmdd) <tag>".format(sys.argv[0]))
        sys.exit(1)
    casting_id = int(sys.argv[1])
    if casting_id not in commoninfo.CastingId2MetaInfo:
        print("CastingId:{} not exist!")
        sys.exit(1)
    casting_name = commoninfo.CastingId2MetaInfo[casting_id][0]
    casting_ch_name = commoninfo.CastingId2MetaInfo[casting_id][1]
    print("获取{}的交易信息...".format(casting_name))
    date = sys.argv[2]
    tag = sys.argv[3]
    if len(tag) == 8:
        # dayly
        start_time = datetime.datetime.strptime(date + " 0:0:0", "%Y%m%d %H:%M:%S")
        end_time = datetime.datetime.strptime(date + " 23:59:59", "%Y%m%d %H:%M:%S")
    elif len(tag) == 10:
        # hourly
        start_time = datetime.datetime.strptime(date + " 0:0:0", "%Y%m%d %H:%M:%S")
        end_time = datetime.datetime.strptime(tag, "%Y%m%d%H")
    else:
        print("错误的tag, tag={}!".format(tag))
        sys.exit(1)
    time_span_str = "时间段：{} - {}".format(start_time.strftime("%Y/%m/%d %H:%M:%S"), end_time.strftime("%Y/%m/%d %H:%M:%S"))

    (res_code, saled_prods) = get_saled_products(casting_id)
    if res_code != 0:
        print("获取在售列表信息失败, res_code={}".format(res_code))
        sys.exit(1)
    print("{} saled info in total.".format(len(saled_prods)))
    
    # 如果已存在交易信息文件，加载
    prodid2detailinfo = {}
    trans_logs_file_name = "data/" + commoninfo.Transaction_Logs_File_Name.format(casting_name)
    if os.path.exists(trans_logs_file_name):
        with open(trans_logs_file_name, encoding="utf-8-sig") as trans_logs_file:
            # 加载已有的交易信息
            for line in trans_logs_file:
                dinfo = [None] * DETAIL_ITEM_COUNT
                items = line.strip().replace(",,", ",").split(",")
                dinfo[DETAIL_SELLER_ID_INDEX] = int(items[DETAIL_SELLER_ID_INDEX])
                dinfo[DETAIL_SELLER_INDEX] = items[DETAIL_SELLER_INDEX]
                dinfo[DETAIL_BUYER_INDEX] = items[DETAIL_BUYER_INDEX]
                dinfo[DETAIL_SALE_TIME_INDEX] = datetime.datetime.strptime(
                    items[DETAIL_SALE_TIME_INDEX], "%Y/%m/%d %H:%M:%S")
                dinfo[DETAIL_PRICE_INDEX] = float(items[DETAIL_PRICE_INDEX])
                dinfo[DETAIL_TOKEN_ID_INDEX] = int(items[DETAIL_TOKEN_ID_INDEX])
                dinfo[DETAIL_PROD_ID_INDEX] = int(items[DETAIL_PROD_ID_INDEX])
                dinfo[DETAIL_DETAIL_ID_INDEX] = int(items[DETAIL_DETAIL_ID_INDEX])
                prodid2detailinfo[dinfo[DETAIL_PROD_ID_INDEX]] = dinfo
    
    scan_cnt = 0
    selluser2cnt = {}
    buyer2cnt = {}
    detail_info_list = []
    trans_cnt = 0
    max_price = None
    min_price = None
    total_price = 0
    for (prod_id, token_id, _) in saled_prods:
        if prod_id in prodid2detailinfo:
            detail_info = prodid2detailinfo[prod_id]
        else:
            detail_info = get_product_detail(prod_id)
            # 防止被封禁
            time.sleep(1)
            if not detail_info:
                print("无法获取ProdId:{}的详细信息!".format(prod_id))
                continue
            detail_info[DETAIL_TOKEN_ID_INDEX] = token_id
            prodid2detailinfo[prod_id] = detail_info
        scan_cnt += 1
        if scan_cnt % 100 == 0:
            print("{} {} saled info scanned.".format(datetime.datetime.now(), scan_cnt))
        
        #print(prod_id, detail_info)
        if detail_info[DETAIL_SALE_TIME_INDEX] < start_time or detail_info[DETAIL_SALE_TIME_INDEX] > end_time:
            # 超出时间范围，忽略
            continue
        trans_cnt += 1
        if max_price == None or detail_info[DETAIL_PRICE_INDEX] > max_price:
            max_price = detail_info[DETAIL_PRICE_INDEX]
        if min_price == None or detail_info[DETAIL_PRICE_INDEX] < min_price:
            min_price = detail_info[DETAIL_PRICE_INDEX]
        total_price += detail_info[DETAIL_PRICE_INDEX]
        detail_info_list.append(detail_info)
        selluser = detail_info[DETAIL_SELLER_ID_INDEX]
        selluser_nickname = detail_info[DETAIL_SELLER_INDEX]
        buyer_nickname = detail_info[DETAIL_BUYER_INDEX]
        if selluser_nickname not in selluser2cnt:
            selluser2cnt[selluser_nickname] = [0, selluser]
        selluser2cnt[selluser_nickname][0] += 1
        if buyer_nickname not in buyer2cnt:
            buyer2cnt[buyer_nickname] = 0
        buyer2cnt[buyer_nickname] += 1
    

    # 刷新交易信息文件
    with open(trans_logs_file_name, "w", encoding="utf-8-sig") as trans_logs_file:
        for (_, dinfo) in prodid2detailinfo.items():
            trans_logs_file.write("{},{},{},{},{},{},{},{}\n".format(
                dinfo[DETAIL_SELLER_ID_INDEX], dinfo[DETAIL_SELLER_INDEX], dinfo[DETAIL_BUYER_INDEX],
                dinfo[DETAIL_SALE_TIME_INDEX].strftime("%Y/%m/%d %H:%M:%S"), dinfo[DETAIL_PRICE_INDEX],
                dinfo[DETAIL_TOKEN_ID_INDEX], dinfo[DETAIL_PROD_ID_INDEX], dinfo[DETAIL_DETAIL_ID_INDEX]
            ))

    # 生成docx文件
    docx_file_name = "data/{}_{}.docx".format(
        casting_name[:2]+casting_ch_name, tag
    )
    doc = Document()


    # 按成交时间倒序排序，输出
    doc.add_heading(casting_ch_name, level=2)
    doc.add_heading("{}交易数据".format(casting_ch_name), level=3)
    doc.add_paragraph(time_span_str)
    detail_info_list.sort(key=lambda dinfo: dinfo[DETAIL_SALE_TIME_INDEX], reverse=True)
    result_file_name = "data/_grab_nft_price_result_{}_{}.csv".format(
        casting_name, tag
    )
    with open(result_file_name, "w", encoding="utf-8-sig") as result_file:
        avg_price = total_price/trans_cnt if trans_cnt > 0 else None 
        if trans_cnt > 0:
            summary = "成交{}笔，最高价格{:.2f}元，最低价格{:.2f}元，均价{:.2f}元。\n".format(
                trans_cnt, max_price, min_price, avg_price
            )
            result_file.write(summary)
            doc.add_paragraph(summary.strip())
        
        table = doc.add_table(1+len(detail_info_list), 6)
        table.style = "TableGrid"
        heading_cells = table.rows[0].cells
        heading_cells[0].text = "卖出者Id"
        heading_cells[1].text = "卖出者昵称"
        heading_cells[2].text = "买入者昵称"
        heading_cells[3].text = "买入时间"
        heading_cells[4].text = "价格"
        heading_cells[5].text = "TokenID"

        result_file.write("{},{},{},{},{},{},{},{}\n".format(
            "卖出者Id", "卖出者昵称", "买入者昵称", "买入时间", "价格", "TokenID", "DEBUG1", "DEBUG2"
        ))
        row_index = 1
        for dinfo in detail_info_list:
            result_file.write("{},{},{},{},{},{},{},{}\n".format(
                dinfo[DETAIL_SELLER_ID_INDEX], dinfo[DETAIL_SELLER_INDEX], dinfo[DETAIL_BUYER_INDEX],
                dinfo[DETAIL_SALE_TIME_INDEX].strftime("%Y/%m/%d %H:%M:%S"), dinfo[DETAIL_PRICE_INDEX],
                "#"+ str(dinfo[DETAIL_TOKEN_ID_INDEX]),
                dinfo[DETAIL_PROD_ID_INDEX], dinfo[DETAIL_DETAIL_ID_INDEX]
            ))
            cells = table.rows[row_index].cells
            cells[0].text = str(dinfo[DETAIL_SELLER_ID_INDEX])
            cells[1].text = dinfo[DETAIL_SELLER_INDEX]
            cells[2].text = dinfo[DETAIL_BUYER_INDEX]
            cells[3].text = dinfo[DETAIL_SALE_TIME_INDEX].strftime("%Y/%m/%d %H:%M:%S")
            cells[4].text = str(dinfo[DETAIL_PRICE_INDEX])
            cells[5].text = "#"+ str(dinfo[DETAIL_TOKEN_ID_INDEX])
            row_index += 1

    # 输出卖出者信息
    doc.add_heading("{}卖出者信息".format(casting_ch_name), level=3)
    doc.add_paragraph(time_span_str)
    sellers_info = []
    for s in selluser2cnt:
        sellers_info.append([s, selluser2cnt[s][1], selluser2cnt[s][0]])
    sellers_info.sort(key=lambda s: s[2], reverse=True)
    result_file_name = "data/_grab_nft_price_result_{}_{}.csv.sellersinfo.csv".format(
        casting_name, tag
    )
    with open(result_file_name, "w", encoding="utf-8-sig") as result_file:
        table = doc.add_table(1+len(sellers_info), 2)
        table.style = "TableGrid"
        heading_cells = table.rows[0].cells
        heading_cells[0].text = "卖出者昵称"
        heading_cells[1].text = "数量"
        row_index = 1
        for (sname, sid, cnt) in sellers_info:
            result_file.write("{},{},{}\n".format(
                sid, sname, cnt
            ))
            cells = table.rows[row_index].cells
            cells[0].text = sname
            cells[1].text = str(cnt)
            row_index += 1
    
    # 输出买入者信息
    doc.add_heading("{}买入者信息".format(casting_ch_name), level=3)
    doc.add_paragraph(time_span_str)
    buyers_info = []
    for b in buyer2cnt:
        buyers_info.append([b, buyer2cnt[b]])
    buyers_info.sort(key=lambda b: b[1], reverse=True)
    result_file_name = "data/_grab_nft_price_result_{}_{}.csv.buyersinfo.csv".format(
        casting_name, tag
    )
    with open(result_file_name, "w", encoding="utf-8-sig") as result_file:
        table = doc.add_table(1+len(buyers_info), 2)
        table.style = "TableGrid"
        heading_cells = table.rows[0].cells
        heading_cells[0].text = "买入者昵称"
        heading_cells[1].text = "数量"
        row_index = 1
        for (bname, cnt) in buyers_info:
            result_file.write("{},{}\n".format(
                bname, cnt
            ))
            cells = table.rows[row_index].cells
            cells[0].text = bname
            cells[1].text = str(cnt)
            row_index += 1
    
    doc.save(docx_file_name)
