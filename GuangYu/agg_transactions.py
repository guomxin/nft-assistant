# coding: utf-8

import sys
import os
import datetime

from docx import Document

from gycommon import commoninfo
from gycommon import utils

#SELLER_NAME_INDEX = 1
#BUYER_NAME_INDEX = 2
TRANS_PRICE_INDEX = 4

def analyze_trans(tag):
    casting2info = {}
    all_total_price = 0.0
    for casting_id in commoninfo.CastingId2MetaInfo:
        casting_name = commoninfo.CastingId2MetaInfo[casting_id][0]
        casting_ch_name = commoninfo.CastingId2MetaInfo[casting_id][1]
        result_file_name = "data/upload/{}/{}/_grab_nft_price_result_{}_{}.csv".format(
            tag, casting_name,
            casting_name, tag
        )
        line_cnt = 0
        trans_min_price = None
        trans_max_price = None
        trans_total_price = 0
        trans_cnt = 0
        with open(result_file_name,  encoding="utf-8-sig") as result_file:
            for line in result_file:
                line_cnt += 1
                if line_cnt <= 2:
                    # 忽略总结行和标题行
                    continue
                items = line.strip().split(",")
                price = float(items[TRANS_PRICE_INDEX])
                all_total_price += price
                trans_cnt += 1
                if not trans_min_price:
                    trans_min_price = price
                if not trans_max_price:
                    trans_max_price = price
                if price < trans_min_price:
                    trans_min_price = price
                if price > trans_max_price:
                    trans_max_price = price
                trans_total_price += price
            casting2info[casting_ch_name] = [trans_cnt, trans_min_price, trans_max_price, 
                trans_total_price / trans_cnt if trans_cnt > 0 else None, trans_total_price]
    
    trans_infos = []
    for casting_ch_name in casting2info:
        info = [casting_ch_name]
        info.extend(casting2info[casting_ch_name])
        trans_infos.append(info)
    trans_infos.sort(key=lambda i: i[-1], reverse=True)
    
    # 生成docx文件
    docx_file_name = "data/日交易额_{}.docx".format(
        tag
    )
    if len(tag) == 8:
        # dayly
        start_time = datetime.datetime.strptime(tag + " 0:0:0", "%Y%m%d %H:%M:%S")
        end_time = datetime.datetime.strptime(tag + " 23:59:59", "%Y%m%d %H:%M:%S")
    elif len(tag) == 10:
        # hourly
        start_time = datetime.datetime.strptime(tag[:8] + " 0:0:0", "%Y%m%d %H:%M:%S")
        end_time = datetime.datetime.strptime(tag, "%Y%m%d%H")
    else:
        print("错误的tag, tag={}!".format(tag))
        return
    time_span_str = "时间段：{} - {}".format(start_time.strftime("%Y/%m/%d %H:%M:%S"), end_time.strftime("%Y/%m/%d %H:%M:%S"))
    doc = Document() # 追加到已有文件
    doc.add_paragraph("{}, 总成交额: {:.2f}万".format(time_span_str, all_total_price/10000))
    table = doc.add_table(1, 4)
    table.style = "TableGrid"
    heading_cells = table.rows[0].cells
    heading_cells[0].text = "名称"
    heading_cells[1].text = "数量"
    heading_cells[2].text = "均价"
    heading_cells[2].text = "合计"

    # 输出
    content = "**{}-总成交额:{:.2f}万**\n".format(tag, all_total_price/10000)
    for (casting_ch_name, cnt, min_price, max_price, avg_price, total_price) in trans_infos:
        if cnt > 0:
            content += "{}\n>数量:{}\n>最低价:{}\n>最高价:{}\n>均价:{:.2f}\n>合计:{:.2f}万\n\n".format(
                casting_ch_name, cnt, min_price,
                max_price, avg_price, total_price/10000)
            cells = table.add_row().cells
            cells[0].text = casting_ch_name
            cells[1].text = str(cnt)
            cells[2].text = "{:.2%}".format(avg_price)
            cells[3].text = "{:.2%}".format(total_price/10000)
    doc.save(docx_file_name)
    
    utils.send_workwx_msg_agg(utils.TradingValue_MSG, "markdown", content)

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("{} <tag>".format(sys.argv[0]))
        sys.exit(1)
    tag = sys.argv[1]
    analyze_trans(tag)

    
